import requests as r
from bs4 import BeautifulSoup as bs
import pandas as pd
import re
from itertools import chain
import matplotlib.pyplot as plt
import numpy as np
import openpyxl
import os
from os.path import exists
import sys
from datetime import datetime, date, timedelta
import xlsxwriter as xs
from openpyxl.workbook.child import INVALID_TITLE_REGEX
from openpyxl import load_workbook
import time
from requests.exceptions import SSLError
from requests.exceptions import ConnectionError
import pytube
from pytube import YouTube
from pytube import Playlist
from pytube import Channel
from youtube_transcript_api import YouTubeTranscriptApi
import nltk
from nltk.corpus import stopwords
import sumy
from sumy.parsers.html import HtmlParser
from sumy.parsers.plaintext import PlaintextParser
from sumy.nlp.tokenizers import Tokenizer
from sumy.summarizers.lsa import LsaSummarizer
from sumy.summarizers.lex_rank import LexRankSummarizer
from sumy.summarizers.luhn import LuhnSummarizer
from sumy.summarizers.reduction import ReductionSummarizer
from sumy.summarizers.text_rank import TextRankSummarizer
from sumy.nlp.stemmers import Stemmer
from sumy.utils import get_stop_words
import ssl
from docx import Document
from docx.shared import Inches
import random

#nltk.download()

###########################################################################################

def parser(url,string="",classe="",parser="html.parser"):
	page=r.get(url)
	parser=bs(page.content,parser)
	parser.encoding = parser.apparent_encoding
	if string=="" :
		return parser
	else:
		if classe=="":
			sous_section=parser.find_all(string)
		else:
			sous_section=parser.find_all(string,{"class": classe})
		return sous_section

def applanir(liste,integer=1):
	i=0
	while i<integer:
		liste=list(chain.from_iterable(liste))
		i=i+1
	return liste

def texte(liste):
	texte=[]
	for i in range(len(liste)):
		lis=liste[i].get_text()
		texte.append(lis)

	texte = " ".join(texte)
	#texte=applanir(texte)
	return texte

def wiki_summary(url,nlines,language,stemmer,case):
	parser_ = parser(url,case)
	texte_wiki=texte(parser_)
	#print(texte_wiki)
	parser_ = PlaintextParser.from_string(texte_wiki, Tokenizer(language))
	stemmer = Stemmer(language)

	summarizer = LuhnSummarizer(stemmer)
	summarizer.stop_words = get_stop_words(language)

	summary=summarizer(parser_.document, nlines)
	summary=pd.DataFrame(summary)
	#print(summary)
	return summary

def word_update(dataframe,name,ope="yes"):

	file=directory_Wikipedia+'\\'+name+'.docx'

	file=file.replace("\\", "\\\\")

	if sys.platform == "win32" and exists(file)!=True:
		document = Document()
	else:
		#document = Document(name+".docx")
		#document.clear_content()
		os.remove(file)
		document = Document()

	document.add_heading(name, 0)


	for i, row in dataframe.iterrows():
		for j, cell in row.items():
			document.add_paragraph(str(cell))


	document.save(file)

	if ope=="yes":
		os.system('"'+file+ '"')


def summarization_wiki(url,stemmer,Language,nlines,name,case):

	sumy_wiki=wiki_summary(url,nlines,Language,stemmer,case)

	word_update(sumy_wiki,name)

###########################################################################################



pd.set_option('display.max_colwidth',100)

ntries=50
nlines=10


####

current_directory = os.getcwd()
parent = os.path.dirname(current_directory)
print(parent)

folder="SummaryWikipedia"
directory_Wikipedia = os.path.join(current_directory, folder)

if not os.path.exists(directory_Wikipedia):
	os.makedirs(directory_Wikipedia)


language_en="english"
stemmer_english = Stemmer(language_en)


url_russia_immigration = "https://en.wikipedia.org/wiki/Russian_emigration_following_the_2022_invasion_of_Ukraine"
name_russia_immigration="Russian_emigration_following_the_2022_invasion_of_Ukraine"


url_russia_ukraine_war_2022 = "https://en.wikipedia.org/wiki/2022_Russian_invasion_of_Ukraine"
name_russia_ukraine_war_2022="2022_Russian_invasion_of_Ukraine"

url_russia_sanction = "https://en.wikipedia.org/wiki/International_sanctions_during_the_2022_Russian_invasion_of_Ukraine"
name_russia_sanction="International_sanctions_during_the_2022_Russian_invasion_of_Ukraine"

url_russia_ecoimpact = "https://en.wikipedia.org/wiki/Economic_impact_of_the_2022_Russian_invasion_of_Ukraine#Russia"
name_russia_ecoimpact="Economic_impact_of_the_2022_Russian_invasion_of_Ukraine#Russia"

url_foreign_aid = "https://en.wikipedia.org/wiki/List_of_foreign_aid_to_Ukraine_during_the_Russo-Ukrainian_War"
name_foreign_aid="List_of_foreign_aid_to_Ukraine_during_the_Russo-Ukrainian_War"


url_gov_react = "https://en.wikipedia.org/wiki/Government_and_intergovernmental_reactions_to_the_2022_Russian_invasion_of_Ukraine"
name_gov_react="Government_and_intergovernmental_reactions_to_the_2022_Russian_invasion_of_Ukraine"

summarization_wiki(url_russia_immigration,stemmer_english,language_en,nlines,name_russia_immigration,"p")
summarization_wiki(url_russia_ukraine_war_2022,stemmer_english,language_en,nlines,name_russia_ukraine_war_2022,"p")
summarization_wiki(url_russia_sanction,stemmer_english,language_en,nlines,name_russia_sanction,"p")
summarization_wiki(url_russia_ecoimpact,stemmer_english,language_en,nlines,name_russia_ecoimpact,"p")
summarization_wiki(url_foreign_aid,stemmer_english,language_en,nlines,name_foreign_aid,"li")
summarization_wiki(url_gov_react,stemmer_english,language_en,nlines,name_gov_react,['p', 'li'])



'''
parser = parser(url_russia_immigration,"p")
texte=texte(parser)
#print(texte)
parser = PlaintextParser.from_string(texte, Tokenizer(language_en))
stemmer = Stemmer(language_en)

summarizer = LuhnSummarizer(stemmer)
summarizer.stop_words = get_stop_words(language_en)

summary=summarizer(parser.document, nlines)
summary=pd.DataFrame(summary)
print(summary)

#for sentence in summarizer(parser.document, nlines):
    #print(sentence)

'''
