import requests as r
from bs4 import BeautifulSoup as bs
import pandas as pd
import re
from itertools import chain
import matplotlib.pyplot as plt
import numpy as np
import openpyxl
import os
from os.path import exists
import sys
from datetime import datetime, date, timedelta
import xlsxwriter as xs
from openpyxl.workbook.child import INVALID_TITLE_REGEX
from openpyxl import load_workbook
import time
from requests.exceptions import SSLError
from requests.exceptions import ConnectionError
import PyPDF2
from PyPDF2 import PdfReader
from PIL import Image 
import PIL 
import tempfile
import io
import urllib3
import time
import selenium
from selenium import webdriver
from selenium.webdriver.common.by import By
import sumy
from sumy.parsers.html import HtmlParser
from sumy.parsers.plaintext import PlaintextParser
from sumy.nlp.tokenizers import Tokenizer
from sumy.summarizers.lsa import LsaSummarizer
from sumy.summarizers.lex_rank import LexRankSummarizer
from sumy.summarizers.luhn import LuhnSummarizer
from sumy.summarizers.reduction import ReductionSummarizer
from sumy.summarizers.text_rank import TextRankSummarizer
from sumy.nlp.stemmers import Stemmer
from sumy.utils import get_stop_words
from docx import Document
from docx.shared import Inches



###########################################################################################

def extract_pdf_text(fichier,number_of_pages=""):
	liste_texte=[]
	reader = PdfReader(directory_pdf+"/"+fichier)

	if number_of_pages=="":
		number_of_pages = len(reader.pages)

	for i in range(number_of_pages):
		page=reader.pages[i]
		text=page.extract_text()
		text = text.replace("\n", " ")
		text = text.replace("\uf0b7", " ")
		liste_texte.append(text)

	return liste_texte


def extract_pdf_image(fichier,number_of_pages=""):
	reader = PdfReader(directory_pdf+"/"+fichier)
	number_of_pages = len(reader.pages)
	count=0
	for i in range(number_of_pages):
		page=reader.pages[i]
		for image_file_object in page.images:
			filepath = os.path.join(directory_pdf, str(count) + image_file_object.name)
			with open(filepath, "wb") as fp:
				fp.write(image_file_object.data)
				count=count+1

def pdf_summary(texte,nlines,language,stemmer):

	parser_ = PlaintextParser.from_string(texte, Tokenizer(language))
	stemmer = Stemmer(language)

	summarizer = LuhnSummarizer(stemmer)
	summarizer.stop_words = get_stop_words(language)

	summary=summarizer(parser_.document, nlines)
	summary=pd.DataFrame(summary)
	#print(summary)
	print (summary)
	return summary


def word_update(dataframe,name,ope="yes"):

	file=directory_pdf+'\\'+name+'.docx'

	file=file.replace("\\", "\\\\")

	if sys.platform == "win32" and exists(file)!=True:
		document = Document()
	else:
		#document = Document(name+".docx")
		#document.clear_content()
		os.remove(file)
		document = Document()

	document.add_heading(name, 0)


	for i, row in dataframe.iterrows():
		for j, cell in row.items():
			document.add_paragraph(str(cell))

	document.save(file)

	if ope=="yes":
		os.system('"'+file+ '"')

def summarization_wiki(texte,stemmer,Language,nlines,name):

	sumy_pdf=pdf_summary(texte,nlines,Language,stemmer)

	word_update(sumy_pdf,name)

###########################################################################################
current_directory = os.getcwd()
parent = os.path.dirname(current_directory)
print(parent)

pd.set_option('display.max_colwidth',100)

folder="pdf"
directory_pdf = os.path.join(current_directory, folder)


pdf_congress=extract_pdf_text("IF12092.pdf",2)
extract_pdf_image("IF12092.pdf")
pdf_congress = ' '.join(pdf_congress)

pdf_house_common=extract_pdf_text("CBP-9481.pdf")
pdf_house_common = ' '.join(pdf_house_common)

pdf_CEPR=extract_pdf_text("172987-global_economic_consequences_of_the_war_in_ukraine_sanctions_supply_chains_and_sustainability.pdf")
pdf_CEPR = ' '.join(pdf_CEPR)

pdf_european_parliement=extract_pdf_text("Rapport_Parlement_Europeen.pdf",13)
pdf_european_parliement = ' '.join(pdf_european_parliement)

pdf_european_EPC=extract_pdf_text("Strengthening_the_impact_of_EU_sanctions_DP.pdf",11)
pdf_european_EPC = ' '.join(pdf_european_EPC)



nlines=6
language_en="english"
stemmer_english = Stemmer(language_en)

summarization_wiki(pdf_congress,stemmer_english,language_en,nlines,"Rapport Congrès Americain")
summarization_wiki(pdf_house_common,stemmer_english,language_en,nlines+4,"Rapport House of Commoon")
summarization_wiki(pdf_CEPR,stemmer_english,language_en,nlines+6,"Rapport CEPR")
summarization_wiki(pdf_european_parliement,stemmer_english,language_en,nlines,"Rapport Parlement Européen")
summarization_wiki(pdf_european_EPC,stemmer_english,language_en,nlines,"Rapport EPC")








