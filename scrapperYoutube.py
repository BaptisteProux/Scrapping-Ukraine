import requests as r
from bs4 import BeautifulSoup as bs
import pandas as pd
import re
from itertools import chain
import matplotlib.pyplot as plt
import numpy as np
import openpyxl
import os
from os.path import exists
import sys
from datetime import datetime, date, timedelta
import xlsxwriter as xs
from openpyxl.workbook.child import INVALID_TITLE_REGEX
from openpyxl import load_workbook
import time
from requests.exceptions import SSLError
from requests.exceptions import ConnectionError
import pytube
from pytube import YouTube
from pytube import Playlist
from pytube import Channel
from youtube_transcript_api import YouTubeTranscriptApi
import nltk
from nltk.corpus import stopwords
import sumy
from sumy.parsers.plaintext import PlaintextParser
from sumy.nlp.tokenizers import Tokenizer
from sumy.summarizers.lsa import LsaSummarizer
from sumy.summarizers.lex_rank import LexRankSummarizer
from sumy.summarizers.luhn import LuhnSummarizer
from sumy.summarizers.reduction import ReductionSummarizer
from sumy.summarizers.text_rank import TextRankSummarizer
from sumy.nlp.stemmers import Stemmer
from sumy.utils import get_stop_words
import ssl
from docx import Document
from docx.shared import Inches
import random
import text2emotion as te

#nltk.download()

###########################################################################################


def punctuator(transcript,langue):
	stop_words = set(stopwords.words(langue))
	transcript = transcript.split(" ")
	transcript = [word if word!="" else " " for word in transcript]

	new_transcript=[]

	for i in range(len(transcript)):
		if i%15 == 0 and transcript[i] not in stop_words :
			new_transcript.append(transcript[i])
			new_transcript.append(".")
			
		#elif i%5 == 0 and i%15 != 0:
			#new_transcript.append(",")
		else:
			new_transcript.append(transcript[i])
			

	#print(transcript_list)
	new_transcript = [word if word!=" " else "" for word in new_transcript]
	new_transcript = " ".join(new_transcript)
	return new_transcript


def get_transcript_playlist(playlist,liste_language):
	texte_transcript_playlist=[]
	title_video=[]
	transcript_playlist=[]
	for url in playlist.video_urls:
		print(url)
		texte_transcript=[]
		try:
			yt=YouTube(url)
			title=yt.title
			title_video.append(title)
			idv=yt.video_id
			transcript=YouTubeTranscriptApi.get_transcript(idv,languages=liste_language)
			for i in range(len(transcript)):
				texte=list(transcript[i].values())
				texte_transcript.append(texte[0])
			texte_transcript = ' '.join(texte_transcript)
			texte_transcript_playlist.append(texte_transcript)
			transcript_playlist.append(texte_transcript_playlist)
			transcript_playlist.append(title_video)
		except:
			print(url+" doesnt have a transcript")
	return transcript_playlist

def summarize_transcript_playlist(transcript_playlist,stemmer,Language,lines,function_summarizer):
	texte_summarize_playlist=[]

	for i in range(len(transcript_playlist)):
		if (Language=="english" and "," not in transcript_playlist[i]) or (Language=="french" and (", " not in transcript_playlist[i] or ". " not in transcript_playlist[i] or transcript_playlist[i][-1]!=".") ):
			transcript_playlist[i]=punctuator(transcript_playlist[i],Language)
		parser = PlaintextParser.from_string(transcript_playlist[i],Tokenizer(Language))
		if Language=="english":
			summarizer = function_summarizer(stemmer)
		else:
			summarizer = function_summarizer()
		summarizer.stop_words = get_stop_words(Language)
		summary=summarizer(parser.document,lines)
		print(summary)
		texte_summarize_playlist.append(summary)
	return texte_summarize_playlist

def multiple_try(parameter1,parameter2,nt,function):
	tries = nt
	for i in range(tries):
		try:
			function(parameter1,parameter2)
		except (SSLError,ConnectionError) :
			if i < tries - 1: # i is zero indexed
				time.sleep(2)
				continue
			else:
				raise
		break

def word_update(dataframe,dftitles,name,liste_emotion,ope="yes"):

	file=directory_Youtube+'\\'+name+'.docx'

	file=file.replace("\\", "\\\\")

	if sys.platform == "win32" and exists(file)!=True:
		document = Document()
	else:
		#document = Document(name+".docx")
		#document.clear_content()
		os.remove(file)
		document = Document()

	document.add_heading(name, 0)


	for i, row in dataframe.iterrows():
		title=dftitles.iloc[i,0]
		p=document.add_paragraph()
		p.add_run(str(title)).bold = True
		for j, cell in row.items():
			paragraph = document.add_paragraph(str(cell))
		p=document.add_paragraph()
		p.add_run("Emotion {}: ".format(i+1))
		p.add_run(str(liste_emotion[i]))



	document.save(file)

	if ope=="yes":
		os.system('"'+file+ '"')

def emotion_dataframe(dataframe):

	liste_emotion=[]

	dataframe = dataframe.apply(lambda x: ''.join(str(x)), axis=1)
	dataframe = dataframe.apply(lambda x: ''.join(x))
	dataframe=pd.DataFrame(dataframe)

	for index, row in dataframe.iterrows():
		emotion_summary=te.get_emotion(row[0])
		liste_emotion.append(emotion_summary)
	
	dataframe_emotion = pd.DataFrame(liste_emotion)
	print(dataframe_emotion)

	return (liste_emotion)

	'''for i, row in dataframe.iterrows():
			for j, sentence in row.items():
				paragraph = str(sentence)'''

	'''text_summary= " ".join(sentence.text for sentence in summary_playlist )
				emotion_summary=te.get_emotion(text_summary)
				print(emotion_summary)'''

def summarization(playlist,liste_langue,stemmer,Language,nlines,name,function_summarizer):
	#liste_langue=['en', 'en-GB']

	#texte_transcript_playlist=multiple_try(playlist_king_generals,liste_langue,ntries,get_transcript_playlist)
	tries = 10
	for i in range(tries):
		try:
			transcript_playlist=get_transcript_playlist(playlist,liste_langue)
		except  :
			if i < tries - 1: # i is zero indexed
				time.sleep(1)
				continue
			else:
				raise
		break

	#transcript_playlist=get_transcript_playlist(playlist,liste_langue)

	#print(texte_transcript_playlist)
	texte_transcript_playlist=transcript_playlist[0]
	#print(texte_transcript_playlist)
	title_transcript_playlist=transcript_playlist[1]
	summary_playlist=summarize_transcript_playlist(texte_transcript_playlist,stemmer,Language,nlines,function_summarizer)
	summary_playlist=pd.DataFrame(summary_playlist)
	print(summary_playlist)
	emotion_summary_playlist=emotion_dataframe(summary_playlist)
	title_transcript_playlist=pd.DataFrame(title_transcript_playlist)
	#print(summary_playlist)
	#print(title_transcript_playlist)
	word_update(summary_playlist,title_transcript_playlist,name,emotion_summary_playlist)




###########################################################################################

playlist_perun=Playlist("https://www.youtube.com/playlist?list=PLbCbj03gdsWwxEZNyy_b0aHKFgmVT3G-3")
playlist_tytelman=Playlist("https://www.youtube.com/playlist?list=PL2zIcogZ92WAjLQ6-bPta3SlXs3w_Icfm")
playlist_king_generals=Playlist("https://www.youtube.com/watch?v=yBZPE9o2gHU&list=PLaBYW76inbX5XWJ8IJd9MKucLnkAzK28E")
playlist_stratpol=Playlist("https://www.youtube.com/watch?v=sfS6VUjAYLM&list=PLmcVqkg1EZfUL5PYlEZNU6RhOxL0fP8Y6")
playlist_euronews_fr=Playlist("https://www.youtube.com/watch?v=KI0oL2jrD8c&list=PLKEy1unUi6hUyib9LM8cbSE4Hrw-545oe")
playlist_conflit_en_carte=Playlist("https://www.youtube.com/watch?v=0NhgscZ7aIc&list=PLmHHUFWJQO3PG8YOcFJ9AcIJLP1NO3ZVP")
playlist_stevius=Playlist("https://www.youtube.com/watch?v=spMa8e5E_Po&list=PLPZ-Y738wLA7WxscRxVVNrHjAUgR5wFNB")

pd.set_option('display.max_colwidth',100)
pd.options.display.max_columns = None

ntries=50
nlines=4


####

current_directory = os.getcwd()
parent = os.path.dirname(current_directory)
print(parent)

folder="transcryptYoutube"
directory_Youtube = os.path.join(current_directory, folder)

if not os.path.exists(directory_Youtube):
	os.makedirs(directory_Youtube)

###

'''
try:
    _create_unverified_https_context = ssl._create_unverified_context
except AttributeError:
    pass
else:
    ssl._create_default_https_context = _create_unverified_https_context

nltk.download()
'''



'''
#texte_transcript_playlist=multiple_try(playlist_king_generals,liste_langue,ntries,get_transcript_playlist)
transcript_playlist_king_general=get_transcript_playlist(playlist_king_generals,liste_langue)
#print(texte_transcript_playlist)
texte_transcript_playlist_king_general=transcript_playlist_king_general[0]
title_transcript_playlist_king_general=transcript_playlist_king_general[1]
summary_playlist_king_general=summarize_transcript_playlist(texte_transcript_playlist_king_general,nlines)
print(summary_playlist_king_general[0])
summary_playlist_king_general=pd.DataFrame(summary_playlist_king_general)
title_transcript_playlist_king_general=pd.DataFrame(title_transcript_playlist_king_general)
print(summary_playlist_king_general)
print(title_transcript_playlist_king_general)
word_update(summary_playlist_king_general,title_transcript_playlist_king_general,"playlist_king_generals_transcript")
'''
liste_langue_en=['en', 'en-GB']
liste_langue_fr=['fr']

Language_english ="english"
Language_french ="french"
stemmer_english = Stemmer(Language_english)

summarization(playlist_king_generals,liste_langue_en,stemmer_english,Language_english,nlines,"playlist_king_generals_transcript",LuhnSummarizer)
summarization(playlist_tytelman,liste_langue_fr,stemmer_english,Language_french,nlines,"playlist_tytelman_transcript",LuhnSummarizer)
summarization(playlist_euronews_fr,liste_langue_fr,stemmer_english,Language_french,nlines,"playlist_euronews_french_transcript",LuhnSummarizer)
summarization(playlist_perun,liste_langue_en,stemmer_english,Language_english,nlines,"playlist_perun_transcript",LuhnSummarizer)
summarization(playlist_conflit_en_carte,liste_langue_fr,stemmer_english,Language_french,nlines,"playlist_conflit_en_carte_transcript",LuhnSummarizer)
summarization(playlist_stratpol,liste_langue_fr,stemmer_english,Language_french,nlines,"playlist_stratpol_transcript",LuhnSummarizer)
summarization(playlist_stevius,liste_langue_fr,stemmer_english,Language_french,nlines,"playlist_stevius",LuhnSummarizer)







